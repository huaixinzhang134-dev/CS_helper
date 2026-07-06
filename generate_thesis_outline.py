#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
生成毕设纲要Word文档
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def create_graduation_thesis_outline():
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 设置段落间距
    style.paragraph_format.space_before = Pt(6)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5
    
    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)
    
    # 标题
    title = doc.add_heading('基于微信小程序的CS电竞助手设计与实现', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0, 51, 102)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('—— 毕业设计纲要')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(102, 102, 102)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    doc.add_paragraph('')  # 空行
    
    # 摘要部分
    add_section_title(doc, '摘要')
    abstract_text = (
        '本课题设计并实现了一款基于微信小程序的CS（Counter-Strike）电竞助手应用，'
        '集成赛事比分查询、选手资料库、猜选手游戏等功能，为CS电竞爱好者提供便捷的信息查询和娱乐互动平台。'
        '系统采用微信小程序原生框架结合TypeScript进行开发，后端依托微信云开发平台实现Serverless架构，'
        '无需独立服务器即可完成数据存储、云函数调用和文件管理等核心功能。'
        '本文档详细阐述了系统的需求分析、架构设计、功能实现以及测试方案，'
        '为后续毕业设计论文的撰写提供完整的纲要指导。'
    )
    add_paragraph(doc, abstract_text)
    
    doc.add_paragraph('')
    
    # 关键词
    keywords = doc.add_paragraph()
    run = keywords.add_run('关键词：')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run = keywords.add_run('微信小程序；CS电竞；云开发；猜选手游戏；Serverless架构')
    run.font.size = Pt(12)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    doc.add_paragraph('')
    
    # 第一章
    add_chapter_title(doc, '第一章 绪论')
    
    add_section_title(doc, '1.1 研究背景与意义')
    add_bullet_points(doc, [
        '电子竞技行业蓬勃发展，CS作为经典FPS游戏拥有庞大的玩家和观众群体',
        'CS电竞粉丝对赛事信息、选手资料、互动娱乐有强烈需求',
        '微信小程序具有即用即走、无需安装的优势，适合轻量级电竞服务场景',
        '现有CS电竞资讯平台多为网页端或独立APP，缺乏便捷的移动端解决方案'
    ])
    
    add_section_title(doc, '1.2 国内外研究现状')
    add_bullet_points(doc, [
        '国外：HLTV.org等网站提供专业CS数据，但移动端体验欠佳',
        '国内：各类电竞APP功能繁杂，缺乏专注于CS的轻量级工具',
        '微信小程序生态中，电竞类小程序多集中于赛事直播，缺少综合性助手工具',
        'Serverless架构在小程序开发中的应用日益成熟'
    ])
    
    add_section_title(doc, '1.3 研究内容与目标')
    add_bullet_points(doc, [
        '构建综合性的CS电竞信息服务平台',
        '实现互动性强的猜选手游戏功能（类似Wordle玩法）',
        '提供实时赛事比分查询服务',
        '探索纯云开发架构在小程序中的应用实践'
    ])
    
    add_section_title(doc, '1.4 论文组织结构')
    add_paragraph(doc, '本文共分为七章，依次为：绪论、相关技术介绍、系统需求分析、系统设计、系统实现、系统测试、总结与展望。各章节循序渐进，从理论基础到实践实现，全面阐述本课题的研究工作。')
    
    # 第二章
    add_chapter_title(doc, '第二章 相关技术介绍')
    
    add_section_title(doc, '2.1 微信小程序框架')
    add_bullet_points(doc, [
        '小程序架构原理：双线程模型（逻辑层与渲染层分离）',
        'WXML（WeiXin Markup Language）：页面结构描述语言',
        'WXSS（WeiXin Style Sheets）：页面样式语言，支持部分CSS特性',
        'JavaScript/TypeScript：逻辑层编程语言',
        '小程序生命周期与页面路由机制'
    ])
    
    add_section_title(doc, '2.2 TypeScript 语言')
    add_bullet_points(doc, [
        'TypeScript类型系统：接口、类型别名、泛型',
        '静态类型检查在大型项目中的优势',
        '在小程序开发中的应用：更好的代码提示和错误检测',
        '编译配置与小程序环境的适配'
    ])
    
    add_section_title(doc, '2.3 微信云开发平台')
    add_bullet_points(doc, [
        '云数据库：文档型NoSQL数据库，支持JSON格式数据存储',
        '云函数：Serverless计算服务，按需执行',
        '云存储：文件上传下载服务，支持图片等资源管理',
        '数据库操作：增删改查、聚合查询、正则表达式搜索',
        '权限控制：小程序端与云端的访问权限管理'
    ])
    
    add_section_title(doc, '2.4 深色模式适配技术')
    add_bullet_points(doc, [
        '系统主题监听机制：wx.onThemeChange API',
        '动态样式切换：根据系统主题动态调整UI颜色',
        'TabBar样式动态更新',
        '用户体验优化：适配不同用户的视觉偏好'
    ])
    
    # 第三章
    add_chapter_title(doc, '第三章 系统需求分析')
    
    add_section_title(doc, '3.1 功能需求分析')
    
    add_subsection_title(doc, '3.1.1 赛事模块')
    add_bullet_points(doc, [
        '实时比分展示：显示正在进行、即将开始、已结束的赛事',
        '赛事状态管理：Live/Upcoming/Finished三种状态',
        '赛事信息展示：赛事名称、参赛队伍、队伍Logo、比分、时间',
        '下拉刷新功能：手动更新赛事数据',
        '智能排序：Live > Upcoming > Finished'
    ])
    
    add_subsection_title(doc, '3.1.2 猜一猜游戏模块')
    add_bullet_points(doc, [
        '个人练习模式：无限次猜测，随机选择目标选手',
        '好友PK模式：邀请好友对战，限定8次尝试机会',
        '猜测反馈机制：正确（绿色）、接近（黄色）、错误（无色）',
        '属性对比：战队、国家、年龄、位置、Major参赛次数',
        '数值方向提示：年龄和Major次数的上下箭头提示',
        '选手搜索：支持模糊搜索、分页加载、防抖优化',
        '分享功能：生成分享卡片邀请好友',
        '广告弹窗机制：模拟视频广告播放'
    ])
    
    add_subsection_title(doc, '3.1.3 资料库模块')
    add_bullet_points(doc, [
        '选手列表展示：分页加载，每页20条记录',
        '模糊搜索：支持姓名、ID搜索，正则表达式匹配',
        '选手详情：展示完整选手信息',
        '图片加载：优先云存储，备用HLTV外链'
    ])
    
    add_subsection_title(doc, '3.1.4 用户模块')
    add_bullet_points(doc, [
        '用户登录/登出功能',
        '登录状态管理：session检查与本地缓存',
        '个人信息展示：昵称、头像、等级、积分'
    ])
    
    add_subsection_title(doc, '3.1.5 管理后台模块')
    add_bullet_points(doc, [
        '选手数据管理：增删改查操作',
        '赛事数据管理：增删改查操作',
        '批量数据导入：云函数分批处理',
        '数据库清空功能'
    ])
    
    add_section_title(doc, '3.2 非功能需求分析')
    add_bullet_points(doc, [
        '性能需求：分页加载优化，突破云数据库100条限制',
        '用户体验：深色模式适配，响应式设计，流畅的交互动画',
        '安全性：云数据库权限控制，防止未授权访问',
        '兼容性：支持不同设备和微信版本'
    ])
    
    add_section_title(doc, '3.3 可行性分析')
    add_bullet_points(doc, [
        '技术可行性：微信小程序生态成熟，云开发平台功能完善',
        '经济可行性：云开发免费额度充足，无需额外服务器成本',
        '操作可行性：界面简洁直观，用户易于上手'
    ])
    
    # 第四章
    add_chapter_title(doc, '第四章 系统设计')
    
    add_section_title(doc, '4.1 系统架构设计')
    add_bullet_points(doc, [
        '整体架构：客户端-云开发平台三层架构',
        '客户端：微信小程序（WXML/WXSS/TypeScript）',
        '云端：云数据库、云函数、云存储',
        '数据流设计：前端请求 → 云数据库查询 → 数据返回 → 页面渲染'
    ])
    
    add_section_title(doc, '4.2 功能模块设计')
    add_bullet_points(doc, [
        '模块划分：赛事模块、游戏模块、资料库模块、用户模块、管理模块',
        '模块职责：各模块独立，通过API服务层统一调用',
        '模块交互：页面组件通过services/api.ts调用云数据库'
    ])
    
    add_section_title(doc, '4.3 数据库设计')
    
    add_subsection_title(doc, '4.3.1 PlayerBase集合（选手信息）')
    add_bullet_points(doc, [
        '_id：云数据库自动生成的唯一标识',
        'playerId：HLTV选手ID',
        'name：选手游戏ID',
        'realName：选手真实姓名',
        'team：当前所属战队',
        'formerTeams：曾经服役的战队列表（数组）',
        'country：国籍',
        'age：年龄',
        'majorAppearances：Major参赛次数',
        'position：位置（狙击手/步枪手/指挥/教练）',
        'avatar：头像路径'
    ])
    
    add_subsection_title(doc, '4.3.2 matches集合（赛事信息）')
    add_bullet_points(doc, [
        '_id：云数据库自动生成的唯一标识',
        'event：赛事名称',
        'status：赛事状态（Live/Upcoming/Finished）',
        'teamA：队伍A信息（name、logo、score）',
        'teamB：队伍B信息（name、logo、score）',
        'time：比赛时间（ISO格式）'
    ])
    
    add_section_title(doc, '4.4 界面设计')
    add_bullet_points(doc, [
        '整体UI风格：暗黑主题电竞风格，深色背景配亮色文字',
        'TabBar导航：四个主页面（赛事、猜一猜、资料库、我的）',
        '页面布局：卡片式设计，信息层次清晰',
        '颜色方案：绿色表示正确，黄色表示接近，无色表示错误'
    ])
    
    add_section_title(doc, '4.5 游戏算法设计')
    
    add_subsection_title(doc, '4.5.1 随机选手选择算法')
    add_bullet_points(doc, [
        '获取选手总数：db.collection.count()',
        '生成随机索引：Math.floor(Math.random() * total)',
        'skip随机位置获取单个选手'
    ])
    
    add_subsection_title(doc, '4.5.2 猜测反馈逻辑')
    add_bullet_points(doc, [
        '战队判断：同一战队=正确；答案选手曾服役于猜测选手战队=接近；否则=错误',
        '国家判断：相同=正确；不同=错误',
        '年龄判断：相同=正确；差值≤2=接近；否则=错误',
        'Major次数判断：相同=正确；差值≤2=接近；否则=错误',
        '位置判断：相同=正确；不同=错误'
    ])
    
    add_subsection_title(doc, '4.5.3 数值方向提示算法')
    add_bullet_points(doc, [
        '年龄方向：猜测值<目标值→向上箭头↑；猜测值>目标值→向下箭头↓',
        'Major次数方向：同上逻辑',
        '帮助用户缩小猜测范围'
    ])
    
    # 第五章
    add_chapter_title(doc, '第五章 系统实现')
    
    add_section_title(doc, '5.1 开发环境搭建')
    add_bullet_points(doc, [
        '微信开发者工具：版本要求与配置',
        'TypeScript编译配置：tsconfig.json设置',
        '云开发环境初始化：wx.cloud.init()调用',
        '项目结构：pages、services、utils、cloudfunctions等目录'
    ])
    
    add_section_title(doc, '5.2 核心功能实现')
    
    add_subsection_title(doc, '5.2.1 赛事模块实现')
    add_bullet_points(doc, [
        '数据获取：matchesCollection.get()获取所有赛事',
        '排序算法：statusOrder映射实现Live>Upcoming>Finished排序',
        '下拉刷新：scroll-view的refresher-enabled属性',
        '状态徽章：动态CSS类名显示不同状态颜色'
    ])
    
    add_subsection_title(doc, '5.2.2 猜一猜游戏实现')
    add_bullet_points(doc, [
        '游戏模式选择：模态弹窗选择个人练习或好友PK',
        '选手搜索：云数据库正则模糊搜索（db.RegExp）',
        '防抖处理：400ms延迟搜索优化用户体验',
        '分页加载：scroll-view的bindscrolltolower事件',
        '猜测处理：processGuess函数实现反馈逻辑',
        '分享功能：onShareAppMessage生成分享卡片',
        'PK房间：本地存储模拟房间管理'
    ])
    
    add_subsection_title(doc, '5.2.3 资料库模块实现')
    add_bullet_points(doc, [
        '分页加载：fetchPlayerListPaginated函数，skip/limit分批查询',
        '总数获取：db.collection.count()获取总记录数',
        'hasMore判断：返回数据量与总数比较',
        '搜索功能：searchPlayers函数支持正则前缀匹配',
        '图片处理：normalizeAvatarUrl函数统一图片路径',
        '页面触底：onReachBottom事件触发加载更多'
    ])
    
    add_subsection_title(doc, '5.2.4 用户模块实现')
    add_bullet_points(doc, [
        '登录流程：wx.getUserProfile获取用户信息',
        '状态检查：wx.checkSession检查session有效性',
        '本地缓存：wx.setStorageSync/wx.getStorageSync',
        '退出登录：清除本地缓存数据'
    ])
    
    add_subsection_title(doc, '5.2.5 管理后台实现')
    add_bullet_points(doc, [
        'CRUD操作：playersCollection.add/update/remove',
        '模态表单：showPlayerModal/showMatchModal控制显示',
        '批量导入：云函数importPlayers分批处理数据',
        '清空功能：遍历删除所有记录'
    ])
    
    add_section_title(doc, '5.3 关键技术实现')
    
    add_subsection_title(doc, '5.3.1 深色模式适配')
    add_bullet_points(doc, [
        '系统主题获取：wx.getSystemInfoSync().theme',
        '主题变化监听：wx.onThemeChange回调',
        'TabBar样式更新：wx.setTabBarStyle动态设置',
        '全局状态管理：app.globalData.isDarkMode'
    ])
    
    add_subsection_title(doc, '5.3.2 分页加载优化')
    add_bullet_points(doc, [
        '云数据库限制：单次查询最多100条',
        '分批策略：skip递增，limit设为100',
        '循环查询：while循环直到返回数据少于请求数量',
        '性能考虑：避免一次性加载大量数据'
    ])
    
    add_subsection_title(doc, '5.3.3 搜索防抖')
    add_bullet_points(doc, [
        '防抖原理：延迟执行，清除前一次定时器',
        '实现方式：setTimeout + clearTimeout',
        '延迟时间：400ms平衡响应速度与性能',
        '新关键词处理：重置页码和搜索结果'
    ])
    
    add_subsection_title(doc, '5.3.4 云函数数据导入')
    add_bullet_points(doc, [
        '分批导入：startIndex和batchSize参数控制',
        '错误处理：try-catch捕获单条导入失败',
        '进度统计：successCount和failCount记录',
        '三种操作：import导入、count计数、clear清空'
    ])
    
    # 第六章
    add_chapter_title(doc, '第六章 系统测试')
    
    add_section_title(doc, '6.1 测试环境')
    add_bullet_points(doc, [
        '测试设备：iOS设备、Android设备、微信开发者工具模拟器',
        '微信版本：8.0及以上',
        '云开发环境：cloud1-2ghbpsm69fa43fcb'
    ])
    
    add_section_title(doc, '6.2 功能测试')
    add_bullet_points(doc, [
        '赛事模块：数据加载、状态显示、下拉刷新',
        '游戏模块：模式选择、猜测逻辑、反馈显示、分享功能',
        '资料库模块：分页加载、搜索功能、详情跳转',
        '用户模块：登录、登出、状态保持',
        '管理模块：增删改查、批量导入'
    ])
    
    add_section_title(doc, '6.3 性能测试')
    add_bullet_points(doc, [
        '数据加载性能：首屏加载时间、分页加载时间',
        '搜索响应时间：防抖延迟、数据库查询时间',
        '内存占用：大量数据加载时的内存使用',
        '网络请求：云数据库调用次数与响应时间'
    ])
    
    add_section_title(doc, '6.4 兼容性测试')
    add_bullet_points(doc, [
        '设备适配：不同屏幕尺寸显示效果',
        '系统版本：iOS和Android兼容性',
        '深色模式：浅色/深色主题切换正常',
        '微信版本：低版本微信的兼容性'
    ])
    
    add_section_title(doc, '6.5 测试结果分析')
    add_paragraph(doc, '总结测试过程中发现的问题及解决方案，评估系统是否满足需求分析中提出的各项指标。')
    
    # 第七章
    add_chapter_title(doc, '第七章 总结与展望')
    
    add_section_title(doc, '7.1 工作总结')
    add_bullet_points(doc, [
        '完成的功能模块：赛事、猜一猜、资料库、用户、管理后台',
        '技术难点与解决方案：分页加载、搜索防抖、深色模式适配',
        '项目特色：纯云开发架构、Wordle式猜选手游戏、好友PK模式'
    ])
    
    add_section_title(doc, '7.2 存在不足')
    add_bullet_points(doc, [
        '数据更新机制：需手动更新赛事和选手数据，缺乏自动化同步',
        '好友PK功能：依赖本地存储，非实时同步，存在数据不一致风险',
        '广告功能：仅为模拟实现，未接入真实广告SDK',
        '用户系统：模拟登录，未实现完整的用户认证体系',
        '数据量限制：云数据库免费额度有限，大数据量场景需优化'
    ])
    
    add_section_title(doc, '7.3 未来展望')
    add_bullet_points(doc, [
        '接入实时赛事数据API：自动同步最新赛事信息',
        '完善好友PK实时对战：使用云开发实时数据推送功能',
        '增加更多游戏模式：团队赛、限时挑战、排行榜等',
        '引入机器学习推荐算法：根据用户猜测习惯推荐选手',
        '社区功能扩展：评论、分享、用户生成内容',
        '多游戏支持：扩展到其他电竞游戏（Dota2、Valorant等）',
        '国际化支持：多语言版本，服务全球CS玩家'
    ])
    
    # 参考文献部分
    add_chapter_title(doc, '参考文献')
    references = [
        '微信小程序官方文档. https://developers.weixin.qq.com/miniprogram/dev/framework/',
        '微信云开发文档. https://developers.weixin.qq.com/miniprogram/dev/wxcloud/basis/getting-started.html',
        'TypeScript官方文档. https://www.typescriptlang.org/docs/',
        '张三. 微信小程序开发实战[M]. 北京: 电子工业出版社, 2022.',
        '李四. Serverless架构与应用实践[M]. 北京: 机械工业出版社, 2023.',
        '王五. 基于微信小程序的电竞资讯平台设计与实现[J]. 计算机应用, 2023, 43(2): 123-128.',
        'HLTV.org. Counter-Strike Statistics. https://www.hltv.org/',
        '赵六. 移动端游戏化学习设计研究[J]. 软件导刊, 2022, 21(5): 45-49.'
    ]
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'[{i}] {ref}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 附录部分
    add_chapter_title(doc, '附录')
    
    add_section_title(doc, '附录A：核心代码片段')
    add_paragraph(doc, '（此处可添加关键功能的代码实现，如猜测反馈算法、分页加载逻辑等）')
    
    add_section_title(doc, '附录B：数据库设计文档')
    add_paragraph(doc, '（此处可添加详细的数据库表结构设计、字段说明、索引设计等）')
    
    add_section_title(doc, '附录C：测试用例表')
    add_paragraph(doc, '（此处可添加完整的测试用例表格，包括测试步骤、预期结果、实际结果等）')
    
    # 致谢
    add_chapter_title(doc, '致谢')
    add_paragraph(doc, '感谢指导老师的悉心指导，感谢同学们在开发过程中提供的宝贵建议，感谢开源社区提供的优秀工具和文档。')
    
    # 保存文档
    output_path = r'C:\Users\50584\Desktop\cs\毕设纲要_517demo.docx'
    doc.save(output_path)
    print(f'文档已生成: {output_path}')

def add_chapter_title(doc, title):
    """添加章标题"""
    heading = doc.add_heading(title, level=1)
    for run in heading.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 51, 102)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    heading.paragraph_format.space_before = Pt(24)
    heading.paragraph_format.space_after = Pt(12)

def add_section_title(doc, title):
    """添加节标题"""
    heading = doc.add_heading(title, level=2)
    for run in heading.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 102, 153)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    heading.paragraph_format.space_before = Pt(18)
    heading.paragraph_format.space_after = Pt(8)

def add_subsection_title(doc, title):
    """添加小节标题"""
    heading = doc.add_heading(title, level=3)
    for run in heading.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(51, 102, 153)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)

def add_paragraph(doc, text):
    """添加普通段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.paragraph_format.line_spacing = 1.5

def add_bullet_points(doc, items):
    """添加项目符号列表"""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(12)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)

if __name__ == '__main__':
    create_graduation_thesis_outline()
